from docx import Document

doc = Document()

# Weekday Schedule
doc.add_heading('Weekday Schedule (Mon–Fri)', level=1)
weekday_table = doc.add_table(rows=1, cols=2)
hdr = weekday_table.rows[0].cells
hdr[0].text, hdr[1].text = 'Time', 'Activity'
weekday_rows = [
    ('07:00–08:00', 'Breakfast & cleanup (1 h)'),
    ('08:00–12:00', 'Study (4 h)'),
    ('12:00–13:00', 'Lunch & break (1 h)'),
    ('13:00–16:00', 'Work block 1 (3 h)'),
    ('16:00–17:00', 'Exercise (M/W/F) or break (T/Th) (1 h)'),
    ('17:00–19:00', 'Work block 2 (2 h)'),
    ('19:00–20:00', 'Dinner & cleanup (1 h)'),
    ('20:00–22:00', 'Leisure / buffer (2 h)'),
    ('22:00–23:00', 'Wind‑down & prep for bed (1 h)'),
]
for t, act in weekday_rows:
    cells = weekday_table.add_row().cells
    cells[0].text, cells[1].text = t, act

doc.add_page_break()

# Weekend Schedule
doc.add_heading('Weekend Schedule', level=1)
weekend_table = doc.add_table(rows=1, cols=3)
hdr = weekend_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Time', 'Saturday', 'Sunday'
weekend_rows = [
    ('07:00–08:00', 'Breakfast & cleanup (1 h)', 'Breakfast & cleanup (1 h)'),
    ('08:00–10:00', 'Cleaning / chores (2 h)', 'Weekly review & planning (2 h)'),
    ('10:00–11:00', 'Exercise (optional 4th session) (1 h)', 'Cleaning / errands (1–2 h)'),
    ('11:00–14:00', 'Work/study catch‑up (2–3 h)', 'Free / social time'),
]
for t, sat, sun in weekend_rows:
    cells = weekend_table.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = t, sat, sun

# Save to the present working directory using a relative path
output_path = './schedule.docx'
doc.save(output_path)
print(f"Schedule saved at {output_path}")
